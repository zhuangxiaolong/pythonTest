
import os
import docx
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from lxml import etree  # 引入lxml解析WebService
#pip install python-docx -U
from docx import Document  # 引入docx生成接口文档
from docx.shared import RGBColor #这个是docx的颜色类
from docx.enum.style import WD_STYLE_TYPE
#pip install translate
from translate import Translator
import requests
from urllib.request import urlopen

def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element and a new w:rPr element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run ()
    r._r.append (hyperlink)

    # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
    # Delete this if using a template that has the hyperlink style in it
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True

    return hyperlink

def decodeHtml(req):
   #print(req.text)
   return req.text

def myTranslator(str):
    #有可能连接不上服务器翻译，重试5次
    tryNum=0
    while tryNum<5:
        tryNum=tryNum+1
        result=doTranslator(str)
        if (result!=str):
            return result
    return str
def doTranslator(str):
    try:
       #准备翻译
       translator= Translator(to_lang="zh")
       list = []
    
       word=""
       for i,s in enumerate(str):
          if s=="_":
             list.append(word)
             word=s
             continue
          if s.isupper():
              if i==0:
                  word=word+s
              else:
                  if word.find("_")!=-1:
                      if len(word)==2:
                         word=s
                         continue
                  list.append(word)
                  word=s  
          else:
             word=word+s
           
           #如果是最后的字符
          if i==len(str)-1:
                  list.append(word)
       result=""
       for w in list:
           #处理特殊字符
           if w=="_":
                      continue
           if w=="s":
                      continue
           if w=="n":
                      continue
           elif w=="_c":
                      result=result+" 编号"
                      continue
           elif w=="_n":
                      result=result+" 名称"
                      continue
           elif w.upper()=="M":
                      continue
           elif w.upper()=="S":
                      continue
           elif w.upper()=="T":
                      continue
           elif w.upper()=="D":
                      continue
           elif w.upper()=="GET":
                      result=result+" 获取"
                      continue
           elif w=="Getdt":
                      result=result+" 获取数据表"
                      continue
           elif w.upper()=="DT":
                      result=result+" 数据表"
                      continue
           elif w=="Dish":
                      result=result+" 菜"
                      continue
           elif w=="Desk":
                      result=result+" 桌"
                      continue
           elif w=="Del":
                      result=result+"删除"
                      continue
           elif w=="Pic":
                      result=result+" 照片"
                      continue
           elif w=="Exec":
                      result=result+" 执行"
                      continue
           elif w=="Depart":
                      result=result+" 部门"
                      continue
           elif w=="Num":
                      result=result+" 数字"
                      continue
           elif w=="Key":
                      result=result+" 密钥"
                      continue
           elif w=="token":
                      result=result+" 密钥"
           elif w=="st":
                      result=result+" 空"
                      continue
           elif w=="Right":
                      result=result+" 权限"
                      continue
           else:
                      result=result+" "+w

       translation = translator.translate(result)
       return translation
    except Exception as e:
       print(e)  # 加入打印
       return str

def BuildWordFile(strFilename,path):
    try:  # 加入try catch
       #准备文档
       filename=strFilename
       document = Document()
       xml = etree.parse(path)
       root = xml.getroot()  # 获取根节点
       list = root.getchildren()
       types = list[0]
       lstSchema = types.getchildren()
       schema = lstSchema[0]
       #准备word数据
       apiCount=1
       #文件标题
       heading=document.add_heading(filename+"接口说明文档", 0)
       #新增样式(第一个参数是样式名称，第二个参数是样式类型：1代表段落；2代表字符；3代表表格)
       style = document.styles.add_style('style name 1',2)
       style.font.color.rgb = RGBColor(165, 27, 71)

       style2 = document.styles.add_style('style name 2',2)
       style2.font.size = 11
       style2.font.name="Arial"
       #参数用的样式
       style3 = document.styles.add_style('style name 3', 3)
       style3.font.size = 11
       style3.font.name="Consolas"
       
       #先生成一个汇总表格
       document.add_paragraph(u'接口调用说明和清单',style='Heading 2')
       
       totalRow=1
       for api in schema.getchildren():
          lstComplexType = api.getchildren()
          complexType = lstComplexType[0]
          lstSequence = complexType.getchildren()
          if len(lstSequence) != 0:
              sequence = lstSequence[0]
              lstElement = sequence.getchildren()
              if len(lstElement) != 0:
                #生成表格
                 totalRow=totalRow+1
                 #break;

       #生成表格
       totalTable = document.add_table(rows=totalRow, cols=3, style='Table Grid')
       totalRowIndex=1
       for api in schema.getchildren():
          lstComplexType = api.getchildren()
          complexType = lstComplexType[0]
          lstSequence = complexType.getchildren()
          if len(lstSequence) != 0:
              sequence = lstSequence[0]
              lstElement = sequence.getchildren()
              if len(lstElement) != 0:
                 apiName=api.get("name", "")
                 totalTable.cell(0, 0).text = "序号"
                 totalTable.cell(0, 1).text = "接口"
                 totalTable.cell(0, 2).text = "说明"
                 totalTable.cell(totalRowIndex, 0).text =str(totalRowIndex)

                 #汇总添加明细的超链接
                 ptotalTableApiName=totalTable.cell(totalRowIndex, 1).paragraphs[0]
                 add_hyperlink(ptotalTableApiName, apiName, '#'+str(totalRowIndex)+"."+apiName)
                 #totalTable.cell(totalRowIndex, 1).text =apiName

                 #翻译接口名字得到接口说明
                 description=myTranslator(apiName)
                 totalTable.cell(totalRowIndex, 2).text =description
                 print(str(totalRowIndex)+"添加汇总："+apiName+" "+description)
                 totalRowIndex=totalRowIndex+1
                 #测试
                 '''
                 if totalRowIndex==5:
                     break
                 '''

       for api in schema.getchildren():
          #获取参数
          lstComplexType = api.getchildren()
          complexType = lstComplexType[0]
          lstSequence = complexType.getchildren()
          if len(lstSequence) != 0:
              sequence = lstSequence[0]
              lstElement = sequence.getchildren()
              if len(lstElement) != 0:
                apiName=api.get("name", "")
                print("接口方法:", apiName)  # 输出节点的标签名
                #添加接口名称
                methodParagraph=document.add_paragraph(str(apiCount)+"."+apiName,style='Heading 3')
                apiCount=apiCount+1
                #run=methodParagraph.add_run()
                #run.font.bold = True
                #run.style='IntenseQuote'
                #增加引用
                #document.add_paragraph('Intense quote', style='Intense Quote')
                
                descriptionParagraph=document.add_paragraph("说明：")
                #翻译接口名字得到接口说明
                description=myTranslator(apiName)
                descriptionParagraph.add_run(description)
                
                requestParagraph=document.add_paragraph("请求")
                #requestParagraph.add_run("请求",style=style2)
                #生成表格
                requestTable = document.add_table(
                          rows=2, cols=2, style='Table Grid')
                requestTable.cell(0, 0).text = "方法"
                requestTable.cell(0, 1).text = "URL"
                requestTable.cell(1, 0).text = "POST"
                requestTable.cell(1, 1).text =apiName
                
                paraParagraph=document.add_paragraph("参数：")
                #paraParagraph.add_run("参数：",style=style2)
  
                #生成表格
                table = document.add_table(
                    rows=len(lstElement)+1, cols=4, style="Table Grid")
                table.cell(0, 0).text = "传入类型"
                table.cell(0, 1).text = "参数"
                table.cell(0, 2).text = "值类型"
                table.cell(0, 3).text = "说明"
                row = 1
                for element in lstElement:
                  print("参数：", element.get("name", ""),
                        element.get("type", ""))
                  table.cell(row, 0).text = "POST"
                  table.cell(row, 1).text = element.get("name", "")
                  table.cell(row, 2).text = element.get("type", "")
                  #翻译参数名
                  paraName=myTranslator(element.get("name", ""))
                  table.cell(row, 3).text = paraName
                  row = row + 1
              returnParagraph=document.add_paragraph("返回：")
              #returnParagraph.add_run("返回：",style=style2)
              #生成结果表格
              responseTable = document.add_table(rows=6, cols=3, style='Table Grid')
              responseTable.cell(0, 0).text = "状态"
              responseTable.cell(0, 1).text = "返回"
              responseTable.cell(0, 2).text = "说明"

              responseTable.cell(1, 0).text = "200"
              responseTable.cell(1, 1).text = "0"
              responseTable.cell(2, 0).text = "403"
              responseTable.cell(3, 0).text = "400"
              responseTable.cell(4, 0).text = "401"
              responseTable.cell(5, 0).text = "500"
              #测试
              '''
              if apiCount==5:
                 break
              '''
          
       document.add_page_break()
       document.save("c:\\"+filename+"接口说明文档.docx")
       print("生成文档完成，"+"c:\\"+filename+"接口说明文档.docx")
    except Exception as e:
       print(e)  # 加入打印


def requestUrl(str_link,filename):
    try:  # 加入try catch
               r = requests.get(str_link)
               html=decodeHtml(r)
               fileFullname="c:/%s.xml"%(filename)
               #先把xml保存下来再去读取
               with open(fileFullname,"w") as f:
                    f.write(html)
               BuildWordFile(filename,fileFullname)

    except Exception as e:
            print(e)  # 加入打印


if __name__ == "__main__":
        URL = "http://localhost:11417/cachewebservices/SwtTotalWebService.asmx?WSDL"
        requestUrl(URL,"SwtTotalWebService")
        print("Done")
